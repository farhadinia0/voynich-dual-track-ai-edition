from __future__ import annotations

import argparse
import csv
import hashlib
import json
import re
import shutil
from collections import Counter, defaultdict
from dataclasses import asdict, dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image


RULESET_VERSION = "pilot-0.1.0"
RESEARCH_DATE = "14 August 2026"

SELECTED_FOLIOS = [
    "f1r",
    "f3r",
    "f31r",
    "f67r1",
    "f70r2",
    "f70v2",
    "f75r",
    "fRos",
    "f88r",
    "f99r",
    "f103r",
    "f116r",
]

IMAGE_PAGE_MAP = {
    "f1r": "page-003.jpg",
    "f3r": "page-007.jpg",
    "f31r": "page-061.jpg",
    "f67r1": "page-121.jpg",
    "f70r2": "page-127.jpg",
    "f70v2": "page-128.jpg",
    "f75r": "page-135.jpg",
    "fRos": "page-158.jpg",
    "f88r": "page-161.jpg",
    "f99r": "page-181.jpg",
    "f103r": "page-183.jpg",
    "f116r": "page-205.jpg",
}

SECTION_NAMES = {
    "T": "Text-only / control",
    "H": "Herbal",
    "A": "Astronomical",
    "Z": "Zodiac",
    "B": "Biological / balneological",
    "C": "Cosmological",
    "P": "Pharmaceutical",
    "S": "Star-marked entries (conventionally called recipes)",
    "?": "Unclassified",
}

PAGE_CONTEXT = {
    "f1r": "A text-dominant opening folio with later marginal additions and the historical Tepenec signature. It is used as a control because illustration-led semantics are unavailable.",
    "f3r": "A herbal folio with a single large plant. A historical annotation proposes a Cretan-dittany-like identification, but the plant identification is not secure.",
    "f31r": "A herbal folio in Currier language B and a different scribal hand from f3r, useful for testing whether rules survive a corpus shift.",
    "f67r1": "An astronomical foldout surface dominated by lunar imagery, concentric circles, radial structure, and circular writing.",
    "f70r2": "A cosmological diagram with a central face, rings of text, and eight surrounding lozenge-like forms.",
    "f70v2": "A zodiac page centered on the paired fish associated with Pisces, surrounded by rings of figures, stars, and labels.",
    "f75r": "A biological or balneological page with two pool-like forms and fourteen human figures, conventionally called nymphs.",
    "fRos": "The large nine-rosette foldout: interconnected medallions, paths, towers, circular writing, and numerous labels.",
    "f88r": "A pharmaceutical page with rows of roots or plant parts, short labels, and paragraph blocks.",
    "f99r": "A pharmaceutical inventory-like page with containers, plant fragments, and many short labels.",
    "f103r": "The opening of the star-marked text section, arranged as compact entries with a seven-pointed star marker.",
    "f116r": "A late star-marked text page near the end of the manuscript; included to test whether the same rule set reaches the corpus boundary.",
}

SOURCE_LINKS = [
    (
        "Yale Beinecke Library: The Beinecke Cipher (Voynich) Manuscript",
        "https://beinecke.library.yale.edu/beinecke/collections/beinecke-cipher-voynich-manuscript",
    ),
    (
        "Yale Digital Collections: Beinecke MS 408",
        "https://collections.library.yale.edu/catalog/2002046",
    ),
    (
        "Rene Zandbergen: Transliteration of the Voynich manuscript text",
        "https://www.voynich.nu/transcr.html",
    ),
    (
        "Zandbergen-Landini EVA transliteration, IVTFF 2.0, version 3b",
        "https://www.voynich.nu/data/ZL3b-n.txt",
    ),
    (
        "Hauer and Kondrak: Decoding Anagrammed Texts Written in an Unknown Language and Script",
        "https://aclanthology.org/Q16-1006/",
    ),
    (
        "Bowern and Lindemann: The Linguistics of the Voynich Manuscript",
        "https://doi.org/10.1146/annurev-linguistics-011619-030613",
    ),
    (
        "Timm and Schinner: A possible generating algorithm of the Voynich manuscript",
        "https://doi.org/10.1080/01611194.2019.1596999",
    ),
    (
        "Greshko: The Naibbe cipher",
        "https://doi.org/10.1080/01611194.2025.2566408",
    ),
]


@dataclass
class LineRecord:
    page: str
    locator: str
    position_code: str
    locus_type: str
    raw: str
    readable: str
    tokens: list[str]


@dataclass
class PageRecord:
    page: str
    meta: dict[str, str]
    comments: list[str]
    lines: list[LineRecord]


def parse_meta(text: str) -> dict[str, str]:
    return dict(re.findall(r"\$([A-Z])=([^\s>]+)", text))


def readable_eva(raw: str) -> str:
    text = raw.strip()
    text = text.replace("<%>", "").replace("<$>", "")
    text = text.replace("<->", " | ")

    def note(match: re.Match[str]) -> str:
        value = match.group(1).strip()
        return f" [note: {value}] " if value else " "

    text = re.sub(r"<!([^>]*)>", note, text)
    text = text.replace(".", " ")
    text = text.replace(",", " [?break] ")
    return re.sub(r"\s+", " ", text).strip()


def normalized_tokens(raw: str) -> list[str]:
    text = raw.lower()
    text = re.sub(r"<![^>]*>", " ", text)
    text = text.replace("<%>", " ").replace("<$>", " ").replace("<->", ".")
    text = re.sub(r"\[([a-z]+):[^\]]+\]", r"\1", text)
    text = re.sub(r"\[\?:[^\]]+\]", " ", text)
    text = text.replace("{", "").replace("}", "")
    text = re.sub(r"@\d+;", "x", text)
    text = text.replace("'", "")
    return [token for token in re.findall(r"[a-z]+", text) if token not in {"note"}]


def locus_type(position_code: str) -> str:
    match = re.search(r"([PLCR])", position_code)
    return match.group(1) if match else "O"


def parse_zl(path: Path) -> list[PageRecord]:
    pages: list[PageRecord] = []
    by_name: dict[str, PageRecord] = {}
    current: PageRecord | None = None

    for source_line in path.read_text(encoding="utf-8", errors="replace").splitlines():
        page_match = re.match(r"^<([^>.]+)>\s*<!\s*(.*?)>\s*$", source_line)
        if page_match:
            name = page_match.group(1)
            current = PageRecord(name, parse_meta(page_match.group(2)), [], [])
            pages.append(current)
            by_name[name] = current
            continue

        if source_line.startswith("#"):
            if current is not None:
                comment = source_line[1:].strip()
                if comment:
                    current.comments.append(comment)
            continue

        line_match = re.match(r"^<([^>]+)>\s*(.*)$", source_line)
        if not line_match or "." not in line_match.group(1):
            continue
        locus_spec = line_match.group(1)
        raw = line_match.group(2).rstrip()
        locator, position = (locus_spec.split(",", 1) + [""])[:2]
        page_name = locator.split(".", 1)[0]
        page = by_name.get(page_name)
        if page is None:
            page = PageRecord(page_name, {}, [], [])
            pages.append(page)
            by_name[page_name] = page
        page.lines.append(
            LineRecord(
                page=page_name,
                locator=locator,
                position_code=position,
                locus_type=locus_type(position),
                raw=raw,
                readable=readable_eva(raw),
                tokens=normalized_tokens(raw),
            )
        )
    return pages


def compute_stats(pages: list[PageRecord]) -> dict:
    all_lines = [line for page in pages for line in page.lines]
    counts = Counter(token for line in all_lines for token in line.tokens)
    initial = Counter(line.tokens[0] for line in all_lines if line.tokens)
    final = Counter(line.tokens[-1] for line in all_lines if line.tokens)
    token_pages: dict[str, set[str]] = defaultdict(set)
    token_sections: dict[str, Counter[str]] = defaultdict(Counter)
    for page in pages:
        code = page.meta.get("I", "?")
        for line in page.lines:
            for token in set(line.tokens):
                token_pages[token].add(page.page)
            for token in line.tokens:
                token_sections[token][code] += 1
    return {
        "all_lines": all_lines,
        "counts": counts,
        "initial": initial,
        "final": final,
        "token_pages": token_pages,
        "token_sections": token_sections,
        "page_total": len(pages),
        "line_total": len(all_lines),
        "token_total": sum(counts.values()),
        "type_total": len(counts),
    }


def token_role(token: str, section: str, stats: dict) -> str:
    count = stats["counts"][token]
    if not count:
        return "unclassified form"
    page_share = len(stats["token_pages"][token]) / max(stats["page_total"], 1)
    init_rate = stats["initial"][token] / count
    final_rate = stats["final"][token] / count
    section_counter = stats["token_sections"][token]
    dominant_section, dominant_n = section_counter.most_common(1)[0]
    dominant_share = dominant_n / count
    if count >= 80 and page_share >= 0.20:
        return "broad high-frequency form"
    if count >= 10 and init_rate >= 0.28:
        return "line-initial tendency"
    if count >= 10 and final_rate >= 0.28:
        return "line-final tendency"
    if count >= 8 and dominant_section == section and dominant_share >= 0.68:
        return "section-concentrated form"
    if token.startswith("qo"):
        return "productive qo- family"
    if token.endswith(("aiin", "ain", "dy", "ey")):
        return "productive ending family"
    return "recurrent form" if count >= 3 else "rare form"


def structural_gloss(line: LineRecord, section: str, stats: dict) -> str:
    n = len(line.tokens)
    location = {
        "P": "running paragraph text",
        "L": "isolated label or short inscription",
        "C": "circular text",
        "R": "radial text",
        "O": "other locus",
    }.get(line.locus_type, "other locus")
    movement = "paragraph opening" if line.position_code.startswith(("@", "*")) else "continuation"
    if "$" in line.raw or line.position_code.startswith("="):
        movement = "paragraph or sequence ending"

    observations = [f"{location}; {movement}; {n} normalized token group{'s' if n != 1 else ''}"]
    if line.tokens:
        candidates: list[str] = []
        for token in line.tokens:
            role = token_role(token, section, stats)
            if role != "rare form":
                candidates.append(f"{token}: {role}")
        unique = list(dict.fromkeys(candidates))[:3]
        if unique:
            observations.append("; ".join(unique))
        repeated = [token for token, count in Counter(line.tokens).items() if count > 1]
        if repeated:
            observations.append("within-line repetition: " + ", ".join(repeated[:3]))
    observations.append("no English lexical value established")
    return ". ".join(observations) + "."


LEXICONS = {
    "T": {
        "verb": ["Record", "Compare", "Preserve", "Mark", "Repeat", "Separate", "Join", "Set down"],
        "object": ["the first account", "the measured portion", "the named substance", "the adjoining entry", "the principal sign", "the prepared matter"],
        "context": ["beside the earlier note", "after the second measure", "under the same sign", "before the final division", "with the remaining portion", "in the appointed order"],
        "close": ["keep the sequence unchanged", "return to the opening mark", "leave the final quantity apart", "compare it with the following entry", "retain only the clear portion"],
    },
    "H": {
        "verb": ["Gather", "Cut", "Bruise", "Dry", "Steep", "Boil", "Strain", "Grind"],
        "object": ["root", "young leaves", "flowering stem", "seed", "outer bark", "fresh shoots", "whole herb"],
        "modifier": ["pale", "bitter", "mature", "tender", "aromatic", "red-veined", "freshly cut"],
        "medium": ["spring water", "wine", "warm oil", "vinegar", "honey", "clear water"],
        "condition": ["until the colour deepens", "in the shade", "over a gentle heat", "before sunrise", "for one measured interval", "until the liquid clears"],
        "purpose": ["for a warming preparation", "for a cooling preparation", "for external use", "as the first ingredient", "for storage in a sealed vessel"],
    },
    "A": {
        "verb": ["Observe", "Mark", "Count", "Follow", "Compare", "Measure"],
        "object": ["the moon's course", "the outer ring", "the returning star", "the central light", "the northern arc", "the interval between signs"],
        "ring": ["the first circle", "the divided band", "the inner course", "the outermost circuit", "the eastern arc"],
        "marker": ["the pale star", "the opposing point", "the marked division", "the lower crossing", "the returning sign"],
        "time": ["before dawn", "at the month's turning", "after the third interval", "when the light increases", "at the closing of the cycle"],
    },
    "Z": {
        "verb": ["mark", "assign", "observe", "number", "place", "compare"],
        "sign": ["the paired fish", "the turning sign", "the watery house", "the double emblem", "the month's threshold"],
        "station": ["the first station", "the outer degree", "the western position", "the returning place", "the lower division"],
        "bearer": ["the star-bearer", "the standing figure", "the vessel-bearer", "the attendant", "the named figure"],
        "marker": ["the blue star", "the outer barrel", "the next division", "the inner ring", "the eastern point"],
    },
    "B": {
        "verb": ["Warm", "Pour", "Draw", "Circulate", "Cool", "Collect"],
        "liquid": ["the green water", "the prepared bath", "the flowing infusion", "the clear liquor", "the heated water"],
        "vessel": ["the upper basin", "the joined vessel", "the lower pool", "the narrow channel", "the covered bath"],
        "subject": ["the bather", "the patient", "the first figure", "the seated woman", "the attendant"],
        "duration": ["for one interval", "until the water cools", "through the second measure", "while the flow remains steady", "until the sign changes"],
        "finish": ["rinse with clear water", "move to the lower basin", "rest before repeating", "close the channel", "retain the remaining liquid"],
    },
    "C": {
        "verb": ["trace", "carry", "join", "measure", "follow", "turn"],
        "center": ["the central enclosure", "the inner rosette", "the marked face", "the middle gate", "the enclosed field"],
        "path": ["the narrow road", "the divided channel", "the outer circuit", "the bridged passage", "the spiral course"],
        "region": ["the eastern tower", "the lower field", "the western circle", "the northern rosette", "the outer wall"],
        "ring": ["the first ring", "the encircling band", "the eightfold division", "the outer medallion", "the returning path"],
        "relation": ["at the marked opening", "without crossing the inner band", "through the narrow gate", "beside the repeated sign", "at the final division"],
    },
    "P": {
        "verb": ["Cut", "Dry", "Grind", "Steep", "Store", "Mix", "Separate", "Bind"],
        "part": ["the root", "the leaf", "the flower head", "the peeled stem", "the seed cluster", "the powdered bark"],
        "ingredient": ["the bitter herb", "the red-rooted plant", "the aromatic plant", "the broad-leaved herb", "the dried fragment"],
        "medium": ["wine", "warm oil", "honey", "vinegar", "clear water", "salted water"],
        "container": ["the tall jar", "the covered vessel", "the narrow-necked pot", "the second container", "the sealed flask"],
        "storage": ["until needed", "away from direct heat", "for one complete cycle", "under a tight cover", "after the moisture has gone"],
    },
    "S": {
        "verb": ["grind", "steep", "boil", "strain", "mix", "warm", "apply"],
        "condition": ["a recurring fever", "pain after exertion", "a cold swelling", "restless sleep", "a dry cough", "weak digestion", "an aching limb"],
        "ingredient": ["the prepared root", "the dried leaves", "the bitter seed", "the aromatic powder", "the clarified resin", "the flower infusion"],
        "medium": ["warm water", "light wine", "honey", "oil", "vinegar", "a clear broth"],
        "time": ["before sleep", "at dawn", "after food", "for three measured days", "while still warm", "once in the evening"],
    },
}


def stable_pick(section: str, category: str, token: str) -> str:
    options = LEXICONS[section][category]
    digest = hashlib.sha256(f"{RULESET_VERSION}|{section}|{category}|{token}".encode("utf-8")).digest()
    return options[int.from_bytes(digest[:4], "big") % len(options)]


def seed_token(tokens: list[str], index: int, fallback: str) -> str:
    return tokens[index % len(tokens)] if tokens else fallback


def speculative_label(line: LineRecord, section: str) -> str:
    tokens = line.tokens
    if section in {"H", "P"}:
        part = stable_pick("P", "part", seed_token(tokens, 0, line.locator))
        ingredient = stable_pick("P", "ingredient", seed_token(tokens, 1, line.locator))
        return f"Label: {part.title()} of {ingredient}."
    if section in {"A", "Z", "C"}:
        center = stable_pick("C", "center", seed_token(tokens, 0, line.locator))
        region = stable_pick("C", "region", seed_token(tokens, 1, line.locator))
        return f"Label: {center.title()}, toward {region}."
    if section == "B":
        vessel = stable_pick("B", "vessel", seed_token(tokens, 0, line.locator))
        return f"Label: {vessel.title()}."
    return "Label: Entry or name marker."


def speculative_sentence(line: LineRecord, section: str) -> str:
    section = section if section in LEXICONS else "T"
    if line.locus_type == "L" or len(line.tokens) <= 2:
        return speculative_label(line, section)
    t = line.tokens
    s = lambda i: seed_token(t, i, line.locator)
    if section == "T":
        return (
            f"{stable_pick(section, 'verb', s(0))} {stable_pick(section, 'object', s(1))} "
            f"{stable_pick(section, 'context', s(2))}; {stable_pick(section, 'close', s(3))}."
        )
    if section == "H":
        return (
            f"{stable_pick(section, 'verb', s(0))} the {stable_pick(section, 'modifier', s(2))} "
            f"{stable_pick(section, 'object', s(1))} with {stable_pick(section, 'medium', s(3))}; "
            f"keep it {stable_pick(section, 'condition', s(4))}, {stable_pick(section, 'purpose', s(5))}."
        )
    if section == "A":
        return (
            f"{stable_pick(section, 'verb', s(0))} {stable_pick(section, 'object', s(1))} along "
            f"{stable_pick(section, 'ring', s(2))}; compare it with {stable_pick(section, 'marker', s(3))} "
            f"{stable_pick(section, 'time', s(4))}."
        )
    if section == "Z":
        return (
            f"Under {stable_pick(section, 'sign', s(0))}, {stable_pick(section, 'verb', s(1))} "
            f"{stable_pick(section, 'station', s(2))}; place {stable_pick(section, 'bearer', s(3))} "
            f"beside {stable_pick(section, 'marker', s(4))}."
        )
    if section == "B":
        return (
            f"{stable_pick(section, 'verb', s(0))} {stable_pick(section, 'liquid', s(1))} in "
            f"{stable_pick(section, 'vessel', s(2))}; let {stable_pick(section, 'subject', s(3))} remain "
            f"{stable_pick(section, 'duration', s(4))}, then {stable_pick(section, 'finish', s(5))}."
        )
    if section == "C":
        return (
            f"From {stable_pick(section, 'center', s(0))}, {stable_pick(section, 'verb', s(1))} "
            f"{stable_pick(section, 'path', s(2))} toward {stable_pick(section, 'region', s(3))}; "
            f"join it to {stable_pick(section, 'ring', s(4))} {stable_pick(section, 'relation', s(5))}."
        )
    if section == "P":
        return (
            f"{stable_pick(section, 'verb', s(0))} {stable_pick(section, 'part', s(1))} of "
            f"{stable_pick(section, 'ingredient', s(2))} with {stable_pick(section, 'medium', s(3))}; "
            f"place it in {stable_pick(section, 'container', s(4))} {stable_pick(section, 'storage', s(5))}."
        )
    return (
        f"For {stable_pick('S', 'condition', s(0))}, take {stable_pick('S', 'ingredient', s(1))}; "
        f"{stable_pick('S', 'verb', s(2))} it with {stable_pick('S', 'medium', s(3))} and use it "
        f"{stable_pick('S', 'time', s(4))}."
    )


def copy_selected_images(image_dir: Path, asset_dir: Path) -> dict[str, Path]:
    asset_dir.mkdir(parents=True, exist_ok=True)
    result: dict[str, Path] = {}
    for folio, filename in IMAGE_PAGE_MAP.items():
        source = image_dir / filename
        if not source.is_file():
            # A repository clone already contains the selected images under
            # their folio names; the page-number names are used only when
            # building directly from the supplied facsimile extraction.
            source = image_dir / f"{folio}.jpg"
        if not source.is_file():
            raise FileNotFoundError(
                f"No image for {folio}: expected {filename} or {folio}.jpg in {image_dir}"
            )
        target = asset_dir / f"{folio}.jpg"
        if source.resolve() != target.resolve():
            shutil.copy2(source, target)
        result[folio] = target
    return result


def export_data(selected: list[PageRecord], stats: dict, data_dir: Path) -> list[dict]:
    data_dir.mkdir(parents=True, exist_ok=True)
    rows: list[dict] = []
    for page in selected:
        section = page.meta.get("I", "?")
        for line in page.lines:
            rows.append(
                {
                    "folio": page.page,
                    "section_code": section,
                    "section": SECTION_NAMES.get(section, section),
                    "locator": line.locator,
                    "position_code": line.position_code,
                    "locus_type": line.locus_type,
                    "raw_eva": line.raw,
                    "display_eva": line.readable,
                    "normalized_tokens": " ".join(line.tokens),
                    "token_count": len(line.tokens),
                    "structural_gloss": structural_gloss(line, section, stats),
                    "structural_confidence": "Moderate for form and position; zero validated English semantics",
                    "speculative_english": speculative_sentence(line, section),
                    "speculative_status": "Creative reconstruction; not a decipherment or translation",
                    "page_context": PAGE_CONTEXT[page.page],
                    "ruleset_version": RULESET_VERSION,
                }
            )

    fieldnames = list(rows[0].keys())
    with (data_dir / "pilot_lines.csv").open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)
    with (data_dir / "pilot_lines.jsonl").open("w", encoding="utf-8") as handle:
        for row in rows:
            handle.write(json.dumps(row, ensure_ascii=False) + "\n")

    token_rows = []
    selected_tokens = Counter(token for page in selected for line in page.lines for token in line.tokens)
    for token, pilot_count in selected_tokens.most_common():
        token_rows.append(
            {
                "token": token,
                "pilot_count": pilot_count,
                "full_corpus_count": stats["counts"][token],
                "full_corpus_pages": len(stats["token_pages"][token]),
                "dominant_section": stats["token_sections"][token].most_common(1)[0][0],
            }
        )
    with (data_dir / "pilot_token_statistics.csv").open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(token_rows[0].keys()))
        writer.writeheader()
        writer.writerows(token_rows)

    ruleset = {
        "ruleset_version": RULESET_VERSION,
        "created": RESEARCH_DATE,
        "selected_folios": SELECTED_FOLIOS,
        "image_page_map": IMAGE_PAGE_MAP,
        "page_context": PAGE_CONTEXT,
        "lexicons": LEXICONS,
        "warning": "The speculative lexicons are deterministic creative devices. They are not decoded Voynich vocabulary.",
    }
    (data_dir / "pilot_ruleset.json").write_text(
        json.dumps(ruleset, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
    )
    return rows


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), name)
    rpr.rFonts.set(qn("w:hAnsi"), name)
    rpr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120):
    if sum(widths_dxa) != 9360:
        raise ValueError(f"Table widths must sum to 9360 DXA, got {sum(widths_dxa)}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    for tag in ("w:tblW", "w:tblInd", "w:tblLayout", "w:tblCellMar", "w:tblBorders"):
        for element in list(tbl_pr.findall(qn(tag))):
            tbl_pr.remove(element)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    margins = OxmlElement("w:tblCellMar")
    for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        margins.append(node)
    tbl_pr.append(margins)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "D7DEE8")
        borders.append(node)
    tbl_pr.append(borders)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def format_table(table, font_size=9.0):
    for row_index, row in enumerate(table.rows):
        if row_index == 0:
            set_repeat_table_header(row)
        for cell in row.cells:
            if row_index == 0:
                set_cell_shading(cell, "E8EEF5")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.08
                for run in paragraph.runs:
                    set_run_font(run, size=font_size, bold=(row_index == 0))


def add_field(paragraph, instruction: str, placeholder: str = "1"):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_hyperlink(paragraph, text: str, url: str):
    relation_id = paragraph.part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E74B5")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.extend([fonts, color, underline])
    run.append(rpr)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def shade_paragraph(paragraph, fill="F4F6F9", border="2E74B5"):
    ppr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    ppr.append(shading)
    p_borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "12")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), border)
    p_borders.append(left)
    ppr.append(p_borders)
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.06)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(10)


def add_callout(doc: Document, label: str, text: str, caution=False):
    paragraph = doc.add_paragraph()
    shade_paragraph(paragraph, fill="FFF8E8" if caution else "F4F6F9", border="A56A00" if caution else "2E74B5")
    run = paragraph.add_run(label + " ")
    set_run_font(run, bold=True, color=(122, 90, 0) if caution else (31, 58, 95))
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def configure_document(doc: Document, running_title: str):
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_specs = {
        "Heading 1": (16, (46, 116, 181), 18, 10),
        "Heading 2": (13, (46, 116, 181), 14, 7),
        "Heading 3": (12, (31, 77, 120), 10, 5),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
        style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(*color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    custom_styles = {
        "EVA Line": ("Courier New", 8.7, (28, 33, 38), 0, 2, 1.0),
        "Analytical Gloss": ("Calibri", 9.2, (80, 88, 96), 0, 5, 1.08),
        "Speculative Line": ("Georgia", 10.2, (35, 50, 62), 0, 6, 1.15),
        "Source Text": ("Calibri", 9.3, (55, 62, 70), 0, 6, 1.10),
    }
    for name, (font, size, color, before, after, spacing) in custom_styles.items():
        style = styles[name] if name in styles else styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = normal
        style.font.name = font
        style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
        style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(*color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = spacing
        if name in {"EVA Line", "Analytical Gloss", "Speculative Line"}:
            style.paragraph_format.keep_together = True

    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(running_title)
    set_run_font(run, size=8.3, color=(99, 107, 115), bold=True)
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run(f"{RULESET_VERSION}  |  Page ")
    set_run_font(run, size=8.3, color=(99, 107, 115))
    add_field(paragraph, "PAGE")


def set_picture_alt_text(inline_shape, title: str, description: str):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def add_folio_image(doc: Document, page: PageRecord, image_path: Path):
    with Image.open(image_path) as image:
        ratio = image.width / image.height
    width = 6.0 if ratio >= 1.35 else 3.15
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(3)
    shape = paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    set_picture_alt_text(
        shape,
        f"Voynich manuscript folio {page.page}",
        f"Facsimile view of Voynich manuscript folio {page.page}. {PAGE_CONTEXT[page.page]}",
    )
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(0)
    caption.paragraph_format.space_after = Pt(8)
    run = caption.add_run(f"Folio {page.page}. Courtesy of the Yale University Library; scan sequence from the supplied facsimile.")
    set_run_font(run, size=8.2, color=(95, 103, 113), italic=True)


def add_cover(doc: Document, title: str, subtitle: str, label: str, cover_image: Path, caution: str):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(48)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(16)
    run = kicker.add_run(label.upper())
    set_run_font(run, size=10.5, color=(165, 106, 0), bold=True)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(9)
    run = paragraph.add_run(title)
    set_run_font(run, size=28, color=(32, 55, 72), bold=True)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(20)
    run = paragraph.add_run(subtitle)
    set_run_font(run, size=14, color=(43, 81, 99), italic=True)

    image_paragraph = doc.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.space_after = Pt(12)
    shape = image_paragraph.add_run().add_picture(str(cover_image), width=Inches(4.4))
    set_picture_alt_text(shape, "Voynich manuscript nine-rosette foldout", PAGE_CONTEXT["fRos"])

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(12)
    run = paragraph.add_run(caution)
    set_run_font(run, size=9.5, color=(155, 28, 28), bold=True)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(f"Pilot edition {RULESET_VERSION} | {RESEARCH_DATE}")
    set_run_font(run, size=9.5, color=(80, 80, 80))
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Prepared with OpenAI Codex - GPT-5.6 Sol (max reasoning)")
    set_run_font(run, size=9.2, color=(80, 80, 80), italic=True)


def add_summary_table(doc: Document, rows: list[tuple[str, str]]):
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Item"
    table.rows[0].cells[1].text = "Pilot result"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    set_table_geometry(table, [2550, 6810])
    format_table(table, font_size=9.2)
    return table


def add_source_list(doc: Document):
    doc.add_page_break()
    doc.add_heading("Sources and citation", level=1)
    doc.add_paragraph(
        "The manuscript images are public-domain historical material held by Yale. The ZL transliteration is a modern scholarly data source and is not relicensed by this edition. The project's original analysis and creative reconstruction are separately licensed in the repository."
    )
    for index, (label, url) in enumerate(SOURCE_LINKS, 1):
        paragraph = doc.add_paragraph(style="Source Text")
        run = paragraph.add_run(f"{index}. ")
        set_run_font(run, size=9.3, bold=True)
        add_hyperlink(paragraph, label, url)
        run = paragraph.add_run(f". {url}")
        set_run_font(run, size=8.2, color=(95, 103, 113))


def build_analytical(
    selected: list[PageRecord], stats: dict, images: dict[str, Path], output: Path
):
    doc = Document()
    configure_document(doc, "Voynich Dual-Track AI Edition | Evidence-Based Pilot")
    add_cover(
        doc,
        "The Voynich Manuscript",
        "Evidence-Based Analytical Pilot",
        "Dual-Track AI Edition",
        images["fRos"],
        "This is not a decipherment or translation. This document contains transliteration and structural analysis only.",
    )

    doc.add_page_break()
    doc.add_heading("Research status and scope", level=1)
    add_callout(
        doc,
        "Bottom line:",
        "The visible script can be transliterated and its distributions can be analysed. No reproducible key currently turns it into validated English. Accordingly, this edition reports every observation at the level actually supported by the data.",
    )
    pilot_lines = sum(len(page.lines) for page in selected)
    pilot_tokens = sum(len(line.tokens) for page in selected for line in page.lines)
    add_summary_table(
        doc,
        [
            ("Primary transliteration", "Zandbergen-Landini EVA, IVTFF 2.0, version 3b (13 May 2025)"),
            ("Pilot coverage", f"{len(selected)} representative folios/surfaces; {pilot_lines:,} loci"),
            ("Pilot normalized token slots", f"{pilot_tokens:,}"),
            ("Full ZL corpus used for statistics", f"{stats['page_total']:,} surfaces; {stats['line_total']:,} loci; {stats['token_total']:,} normalized token slots"),
            ("Semantic claim", "None. English word meanings remain unestablished."),
            ("Reproducibility", f"All generated rows identify source locator and ruleset {RULESET_VERSION}."),
        ],
    )

    doc.add_heading("Why two editions", level=2)
    doc.add_paragraph(
        "This analytical volume protects the distinction between evidence and imagination. Its companion volume deliberately explores what a readable medieval herbal, cosmological, bathing, pharmaceutical, and recipe-like reconstruction might sound like. The companion is useful as a creative hypothesis generator, but it is not evidence for what the manuscript says."
    )
    doc.add_heading("Frozen analytical rules", level=2)
    doc.add_paragraph(
        "The parser preserves folio and locus identifiers, keeps uncertain breaks explicit, records whether text is paragraphal, circular, radial, or label-like, and computes token frequency and positional tendencies from the complete ZL corpus. A structural gloss may describe these features, but it may not assign an English dictionary meaning."
    )

    doc.add_heading("Competing models retained", level=2)
    table = doc.add_table(rows=1, cols=3)
    for index, value in enumerate(("Model", "What it can explain", "What remains unproved")):
        table.rows[0].cells[index].text = value
    hypothesis_rows = [
        ("Unknown natural language", "Section clustering, repeated morphology-like families, topic-related distribution", "Language identity, phonetic values, grammar, and translation"),
        ("Ciphertext or codebook", "Patterned output may conceal Latin, Italian, or another plaintext; verbose ciphers can imitate Voynich statistics", "The actual historical key and a unique reverse mapping"),
        ("Generated pseudo-text", "Local similarity, repetition, and copy-with-variation effects", "Whether all long-range and section-related structure follows from generation alone"),
        ("Published language claims", "Individual proposals can fit selected readings or illustrations", "Stable full-corpus rules, held-out prediction, and independent specialist validation"),
    ]
    for row_values in hypothesis_rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            cells[index].text = value
    set_table_geometry(table, [2250, 3500, 3610])
    format_table(table, font_size=8.8)

    doc.add_page_break()
    doc.add_heading("Line-by-line analytical edition", level=1)
    add_callout(
        doc,
        "How to read:",
        "EVA is a visual transliteration alphabet. [?break] marks an uncertain separation and | marks an illustration or major interruption. The gray line beneath each EVA locus is a structural description, not a translation.",
    )
    for page in selected:
        doc.add_page_break()
        section = page.meta.get("I", "?")
        heading = doc.add_heading(f"Folio {page.page} | {SECTION_NAMES.get(section, section)}", level=1)
        heading.paragraph_format.keep_with_next = True
        doc.add_paragraph(PAGE_CONTEXT[page.page])
        add_folio_image(doc, page, images[page.page])
        metadata = []
        if page.meta.get("L"):
            metadata.append(f"Currier language {page.meta['L']}")
        if page.meta.get("H"):
            metadata.append(f"hand {page.meta['H']}")
        if page.meta.get("Q"):
            metadata.append(f"quire {page.meta['Q']}")
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(8)
        run = paragraph.add_run("Metadata: " + ("; ".join(metadata) if metadata else "not assigned"))
        set_run_font(run, size=9.0, color=(95, 103, 113), italic=True)
        for line in page.lines:
            paragraph = doc.add_paragraph(style="EVA Line")
            paragraph.paragraph_format.keep_with_next = True
            run = paragraph.add_run(f"[{line.locator} {line.position_code}] ")
            set_run_font(run, name="Courier New", size=7.8, color=(95, 103, 113), bold=True)
            run = paragraph.add_run(line.readable or "[empty or unreadable locus]")
            set_run_font(run, name="Courier New", size=8.7)
            paragraph = doc.add_paragraph(structural_gloss(line, section, stats), style="Analytical Gloss")
            paragraph.paragraph_format.widow_control = False

    add_source_list(doc)
    add_callout(
        doc,
        "Citation warning:",
        "Do not cite any structural gloss as an English translation. Cite it only as an observation produced by this versioned analytical pipeline.",
        caution=True,
    )
    core = doc.core_properties
    core.title = "The Voynich Manuscript: Evidence-Based Analytical Pilot"
    core.subject = "Pilot transliteration and structural analysis for a dual-track AI edition"
    core.author = ""
    core.last_modified_by = ""
    core.keywords = "Voynich Manuscript, MS 408, EVA, transliteration, structural analysis"
    core.comments = "No semantic decipherment is claimed."
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def build_speculative(selected: list[PageRecord], images: dict[str, Path], output: Path):
    doc = Document()
    configure_document(doc, "Voynich Dual-Track AI Edition | Speculative Reconstruction")
    add_cover(
        doc,
        "The Voynich Manuscript",
        "Speculative English Reconstruction Pilot",
        "Dual-Track AI Edition",
        images["fRos"],
        "NOT A TRANSLATION. Creative reconstruction only - not a decipherment or statement of historical fact.",
    )

    doc.add_page_break()
    doc.add_heading("What this volume is", level=1)
    add_callout(
        doc,
        "Mandatory label:",
        "Every English line in this document is generated from a frozen, genre-constrained literary rule. No Voynich word has been shown to mean the English word placed beside it.",
        caution=True,
    )
    doc.add_paragraph(
        "The reconstruction asks a narrow creative question: if the illustrations broadly indicate herbal, astronomical, zodiacal, bathing, cosmological, pharmaceutical, and recipe-like genres, what kind of concise medieval technical prose could occupy the same line structure? It does not answer the historical question of what the unknown author actually wrote."
    )
    doc.add_heading("Generation rule", level=2)
    doc.add_paragraph(
        "For each locus, the script keeps the source folio, locus type, token count, and EVA text. A cryptographic hash of the ruleset version, manuscript section, semantic slot, and EVA token selects words from a fixed section-specific English vocabulary. This makes the output deterministic and auditable while demonstrating that consistency alone does not prove decipherment."
    )
    doc.add_heading("Permitted and prohibited use", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Permitted description"
    table.rows[0].cells[1].text = "Misleading description"
    for left, right in (
        ("Speculative AI reconstruction", "English translation of the Voynich Manuscript"),
        ("Genre-constrained literary experiment", "Decoded medieval recipe"),
        ("Deterministic hypothesis generator", "Evidence that a particular EVA token has this meaning"),
    ):
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
    set_table_geometry(table, [4680, 4680])
    format_table(table, font_size=9.2)

    doc.add_page_break()
    doc.add_heading("Line-by-line speculative reconstruction", level=1)
    for page in selected:
        doc.add_page_break()
        section = page.meta.get("I", "?")
        heading = doc.add_heading(f"Folio {page.page} | {SECTION_NAMES.get(section, section)}", level=1)
        heading.paragraph_format.keep_with_next = True
        doc.add_paragraph(PAGE_CONTEXT[page.page])
        add_folio_image(doc, page, images[page.page])
        for line in page.lines:
            paragraph = doc.add_paragraph(style="EVA Line")
            paragraph.paragraph_format.keep_with_next = True
            run = paragraph.add_run(f"[{line.locator} {line.position_code}] ")
            set_run_font(run, name="Courier New", size=7.8, color=(95, 103, 113), bold=True)
            run = paragraph.add_run(line.readable or "[empty or unreadable locus]")
            set_run_font(run, name="Courier New", size=8.7)
            paragraph = doc.add_paragraph(style="Speculative Line")
            run = paragraph.add_run("Speculative English: ")
            set_run_font(run, name="Georgia", size=9.2, color=(155, 28, 28), bold=True, italic=False)
            run = paragraph.add_run(speculative_sentence(line, section))
            set_run_font(run, name="Georgia", size=10.2, color=(35, 50, 62), italic=True)

    add_source_list(doc)
    add_callout(
        doc,
        "Final warning:",
        "The English prose above was invented by a declared algorithm. It must never be quoted as the manuscript's recovered plaintext.",
        caution=True,
    )
    core = doc.core_properties
    core.title = "The Voynich Manuscript: Speculative English Reconstruction Pilot"
    core.subject = "Deterministic creative companion to an evidence-based analytical edition"
    core.author = ""
    core.last_modified_by = ""
    core.keywords = "Voynich Manuscript, speculative reconstruction, AI, creative edition"
    core.comments = "Not a decipherment or verified translation."
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def write_manifest(repo: Path, source: Path, selected: list[PageRecord], rows: list[dict]):
    payload = {
        "project": "Voynich Dual-Track AI Edition",
        "ruleset": RULESET_VERSION,
        "generated": RESEARCH_DATE,
        "source_file": source.name,
        "source_sha256": hashlib.sha256(source.read_bytes()).hexdigest(),
        "selected_folios": SELECTED_FOLIOS,
        "pilot_loci": len(rows),
        "pilot_tokens": sum(row["token_count"] for row in rows),
        "folio_loci": {page.page: len(page.lines) for page in selected},
        "ai_system": "OpenAI Codex",
        "ai_model": "GPT-5.6 Sol",
        "reasoning_effort": "max",
        "disclaimer": "The analytical edition contains no semantic decipherment. The speculative edition is creative reconstruction only.",
    }
    (repo / "data" / "build_manifest.json").write_text(
        json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
    )


def main():
    parser = argparse.ArgumentParser(
        description="Build the two-document Voynich pilot and its auditable data exports."
    )
    parser.add_argument("--zl", type=Path, required=True, help="ZL3b-n IVTFF source file")
    parser.add_argument(
        "--image-dir",
        type=Path,
        required=True,
        help="Directory containing page-NNN.jpg facsimile pages or selected folio-named JPGs",
    )
    parser.add_argument("--repo", type=Path, required=True, help="Repository root")
    args = parser.parse_args()

    pages = parse_zl(args.zl)
    by_name = {page.page: page for page in pages}
    missing = [folio for folio in SELECTED_FOLIOS if folio not in by_name]
    if missing:
        raise ValueError(f"Selected folios missing from source: {missing}")
    selected = [by_name[folio] for folio in SELECTED_FOLIOS]
    stats = compute_stats(pages)
    images = copy_selected_images(args.image_dir, args.repo / "assets" / "folios")
    rows = export_data(selected, stats, args.repo / "data")
    write_manifest(args.repo, args.zl, selected, rows)

    analytical = args.repo / "editions" / "Voynich_Pilot_Evidence_Based_Analytical_Edition.docx"
    speculative = args.repo / "editions" / "Voynich_Pilot_Speculative_English_Reconstruction.docx"
    build_analytical(selected, stats, images, analytical)
    build_speculative(selected, images, speculative)
    print(json.dumps({
        "analytical": str(analytical),
        "speculative": str(speculative),
        "folios": len(selected),
        "loci": len(rows),
        "tokens": sum(row["token_count"] for row in rows),
        "corpus_loci": stats["line_total"],
    }, indent=2))


if __name__ == "__main__":
    main()
