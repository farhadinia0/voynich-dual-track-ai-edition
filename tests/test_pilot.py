from __future__ import annotations

import csv
import json
import unittest
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
EXPECTED_FOLIOS = {
    "f1r", "f3r", "f31r", "f67r1", "f70r2", "f70v2",
    "f75r", "fRos", "f88r", "f99r", "f103r", "f116r",
}
EXPECTED_SOURCE_SHA256 = "bf5b6d4ac1e3a51b1847a9c388318d609020441ccd56984c901c32b09beccafc"


class PilotDataTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        with (ROOT / "data" / "pilot_lines.csv").open(encoding="utf-8", newline="") as handle:
            cls.rows = list(csv.DictReader(handle))
        cls.manifest = json.loads((ROOT / "data" / "build_manifest.json").read_text(encoding="utf-8"))

    def test_expected_scope(self):
        self.assertEqual(len(self.rows), 535)
        self.assertEqual({row["folio"] for row in self.rows}, EXPECTED_FOLIOS)
        self.assertEqual(sum(int(row["token_count"]) for row in self.rows), 3380)

    def test_provenance_fields_are_present(self):
        for row in self.rows:
            self.assertTrue(row["locator"])
            self.assertTrue(row["raw_eva"])
            self.assertTrue(row["display_eva"])
            self.assertEqual(row["ruleset_version"], "pilot-0.1.0")

    def test_claim_boundaries_are_machine_readable(self):
        for row in self.rows:
            self.assertIn("zero validated English semantics", row["structural_confidence"])
            self.assertEqual(
                row["speculative_status"],
                "Creative reconstruction; not a decipherment or translation",
            )

    def test_manifest(self):
        self.assertEqual(self.manifest["source_sha256"], EXPECTED_SOURCE_SHA256)
        self.assertEqual(self.manifest["pilot_loci"], 535)
        self.assertEqual(self.manifest["pilot_tokens"], 3380)
        self.assertEqual(self.manifest["ai_system"], "OpenAI Codex")
        self.assertEqual(self.manifest["ai_model"], "GPT-5.6 Sol")
        self.assertEqual(self.manifest["reasoning_effort"], "max")


class PilotDocumentTests(unittest.TestCase):
    def document_text(self, filename: str) -> str:
        doc = Document(ROOT / "editions" / filename)
        parts = [paragraph.text for paragraph in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        return "\n".join(parts)

    def test_analytical_disclaimer(self):
        text = self.document_text("Voynich_Pilot_Evidence_Based_Analytical_Edition.docx")
        self.assertIn("This is not a decipherment", text)
        self.assertIn("No reproducible key currently turns it into validated English", text)
        self.assertNotIn("Speculative English:", text)

    def test_speculative_disclaimer_and_all_folios(self):
        text = self.document_text("Voynich_Pilot_Speculative_English_Reconstruction.docx")
        self.assertIn("NOT A TRANSLATION", text)
        self.assertIn("Creative reconstruction", text)
        for folio in EXPECTED_FOLIOS:
            self.assertIn(f"Folio {folio}", text)

    def test_ai_disclosure_and_scrubbed_authorship_metadata(self):
        filenames = (
            "Voynich_Pilot_Evidence_Based_Analytical_Edition.docx",
            "Voynich_Pilot_Speculative_English_Reconstruction.docx",
        )
        for filename in filenames:
            with self.subTest(filename=filename):
                text = self.document_text(filename)
                self.assertIn(
                    "Prepared with OpenAI Codex - GPT-5.6 Sol (max reasoning)",
                    text,
                )
                doc = Document(ROOT / "editions" / filename)
                self.assertFalse(doc.core_properties.author)
                self.assertFalse(doc.core_properties.last_modified_by)


if __name__ == "__main__":
    unittest.main()
